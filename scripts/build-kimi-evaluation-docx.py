from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "kimi-evaluation" / "latest"
RESULTS_JSON = REPORT_DIR / "evaluation-results.json"
OUTPUT_DOCX = REPORT_DIR / "kimi-evaluation-questions-and-answers.docx"


def main() -> None:
    data = json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    summary = data["summary"]
    results = data["results"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    configure_styles(doc)
    add_title(doc, summary)
    add_summary_table(doc, summary)
    add_note(doc)

    doc.add_heading("测试问题与回答记录", level=1)
    for index, item in enumerate(results, start=1):
        if index > 1:
            doc.add_page_break()
        add_case(doc, item)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(23, 32, 51)

    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def add_title(doc: Document, summary: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KIMI 学习助教评测问答记录")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(23, 32, 51)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"模型：{summary.get('model')}  |  Provider：{summary.get('provider')}  |  用例数：{summary.get('caseCount')}"
    )
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor(90, 102, 120)


def add_summary_table(doc: Document, summary: dict) -> None:
    doc.add_heading("评测摘要", level=1)
    rows = [
        ("生成时间", summary.get("createdAt")),
        ("真实模型回答数", summary.get("realLlmCount")),
        ("本地门控/模板回答数", summary.get("templateFallbackCount")),
        ("低置信度回答数", summary.get("lowConfidenceCount")),
        ("知识库检索调用数", summary.get("retrievalCalledCount")),
        ("拒绝编造次数", summary.get("refusalCount")),
        ("API Key 是否写入报告", summary.get("keySaved")),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [5.0, 10.5])
    set_cell_text(table.rows[0].cells[0], "项目", bold=True, fill="DDEAF3")
    set_cell_text(table.rows[0].cells[1], "结果", bold=True, fill="DDEAF3")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], str(label), bold=True)
        set_cell_text(cells[1], str(value))


def add_note(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "说明：本文档整理最新一轮本地评测中的问题、回答、生成模式、置信度、answerability 与引用来源。"
        "其中 template_fallback 表示 agent 在证据不足或拒绝编造场景下使用本地安全门控回答，不代表真实模型回答。"
    )
    run.font.color.rgb = RGBColor(90, 102, 120)


def add_case(doc: Document, item: dict) -> None:
    response = item["response"]
    trace = response["decisionTrace"]
    doc.add_heading(f"{item['id']}  {item['title']}", level=2)

    meta = doc.add_table(rows=1, cols=4)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(meta, [3.7, 3.7, 4.2, 3.9])
    values = [
        ("mode", response.get("answerGenerationMode")),
        ("confidence", response.get("confidence")),
        ("answerability", trace.get("answerability", {}).get("status")),
        ("skill", (response.get("usedSkills") or [{}])[0].get("status", "none")),
    ]
    for cell, (label, value) in zip(meta.rows[0].cells, values):
        set_cell_text(cell, f"{label}\n{value}", bold=True, fill="F4F8FB")

    doc.add_heading("问题", level=3)
    add_boxed_paragraph(doc, item["query"], "EAF2F8")

    add_material_preview(doc, item)

    doc.add_heading("回答", level=3)
    add_answer_paragraphs(doc, response.get("answer", ""))

    citations = response.get("citations") or []
    doc.add_heading("引用来源", level=3)
    if citations:
        for citation in citations:
            text = citation.get("title") or citation.get("sourceId") or citation.get("sectionTitle") or "untitled"
            doc.add_paragraph(f"{citation.get('sourceType')}: {text}", style=None)
    else:
        doc.add_paragraph("无。")


def add_material_preview(doc: Document, item: dict) -> None:
    material = item.get("material") or {}
    doc.add_heading("当前 PPT 页", level=3)
    caption = f"{material.get('fileName') or material.get('id') or 'material'} / page {material.get('pageIndex')} / {material.get('pageTitle') or 'untitled'}"
    p = doc.add_paragraph(caption)
    p.paragraph_format.space_after = Pt(6)
    preview_path = material.get("previewImagePath")
    if preview_path and Path(preview_path).exists():
        doc.add_picture(str(preview_path), width=Cm(14.5))
    else:
        doc.add_paragraph(f"预览不可用：{material.get('previewStatus') or 'unavailable'} {material.get('previewError') or ''}".strip())


def add_answer_paragraphs(doc: Document, answer: str) -> None:
    parts = [part.rstrip() for part in answer.splitlines()]
    for part in parts:
        if not part.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(part)


def add_boxed_paragraph(doc: Document, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [15.5])
    set_cell_text(table.rows[0].cells[0], text, fill=fill)


def set_table_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    if fill:
        set_cell_fill(cell, fill)
    set_cell_margins(cell, top=90, start=120, bottom=90, end=120)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


if __name__ == "__main__":
    main()
