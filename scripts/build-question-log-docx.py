from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "evaluation" / "latest"
RESULTS_PATH = REPORT_DIR / "evaluation-results.json"
OUTPUT_PATH = REPORT_DIR / "learning-assistant-test-questions.docx"


def main() -> None:
    payload = json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    summary = payload["summary"]
    results = payload["results"]

    doc = Document()
    setup_document(doc)
    add_title(doc)
    add_meta(doc, summary)
    add_summary_table(doc, summary)
    add_intro(doc)
    add_question_table(doc, results)
    add_detail_sections(doc, results)
    add_rerun_notes(doc)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 22, RGBColor(23, 32, 51)),
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(31, 78, 121)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Learning Assistant Agent 测试问题清单")
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("用于查看测试问题、压力点、agent 决策摘要和结果。")
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor(83, 96, 121)


def add_meta(doc: Document, summary: dict) -> None:
    generated = summary.get("generatedAt", "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"生成时间：{generated}    "
        f"评测通过：{summary.get('passedEvaluationCases')}/{summary.get('totalEvaluationCases')}    "
        f"Node 测试：{'通过' if summary.get('nodeTestPassed') else '失败'}"
    )
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(83, 96, 121)


def add_summary_table(doc: Document, summary: dict) -> None:
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["评测场景", "通过场景", "Node 测试", "外部 API 调用"]
    values = [
        str(summary.get("totalEvaluationCases")),
        str(summary.get("passedEvaluationCases")),
        "通过" if summary.get("nodeTestPassed") else "失败",
        str(summary.get("externalApiCalls", 0)),
    ]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, center=True)
    for i, value in enumerate(values):
        set_cell_text(table.cell(1, i), value, center=True)
    doc.add_paragraph()


def add_intro(doc: Document) -> None:
    doc.add_heading("阅读说明", level=1)
    lines = [
        "这份文档整理自动评测中实际问过的问题，以及每个问题要验证的能力点。",
        "完整机器可读结果见 reports/evaluation/latest/evaluation-results.json；文字日志见 evaluation-log.md。",
        "当前测试不依赖外部模型 API；LLM provider 通过可注入接口验证，后续可以替换为真实模型。",
    ]
    for line in lines:
        doc.add_paragraph(line)


def add_question_table(doc: Document, results: list[dict]) -> None:
    doc.add_heading("问题总表", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    widths = [1.1, 1.4, 1.5, 6.2, 2.0, 1.5, 1.5]
    headers = ["ID", "压力等级", "结果", "测试问题", "意图", "Skill", "引用数"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "1F4E79")
        set_cell_text(cell, header, bold=True, center=True, color=RGBColor(255, 255, 255))

    for item in results:
        cells = table.add_row().cells
        values = [
            item["id"],
            item["pressure"],
            "PASS" if item["passed"] else "FAIL",
            item["query"],
            item.get("decisionTrace", {}).get("detectedIntent", ""),
            (item.get("usedSkills") or [{}])[0].get("status", "none"),
            str(len(item.get("citations") or [])),
        ]
        for idx, value in enumerate(values):
            set_cell_width(cells[idx], widths[idx])
            set_cell_text(cells[idx], str(value), center=idx in [0, 1, 2, 4, 5, 6])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def add_detail_sections(doc: Document, results: list[dict]) -> None:
    doc.add_heading("逐项测试记录", level=1)
    for item in results:
        doc.add_heading(f"{item['id']}  {item['title']}", level=2)
        add_kv(doc, "测试问题", item["query"])
        add_kv(doc, "压力等级", item["pressure"])
        add_kv(doc, "通过结果", "PASS" if item["passed"] else "FAIL")
        add_kv(doc, "预期检查点", "；".join(item.get("expected", [])))
        add_kv(doc, "决策摘要", format_trace(item.get("decisionTrace") or {}))
        add_kv(doc, "Skill 调用", format_skill(item.get("usedSkills") or []))
        add_kv(doc, "教学策略", format_policy(item.get("teachingPolicy") or {}))
        add_kv(doc, "回答摘要", item.get("answerPreview", ""))


def add_rerun_notes(doc: Document) -> None:
    doc.add_heading("复跑命令", level=1)
    commands = [
        'cd /d "E:\\Desktop\\TASKS AND WORK\\问答agent"',
        "npm.cmd run test",
        "npm.cmd run evaluate",
        "npm.cmd run demo",
        "npm.cmd run ui",
    ]
    for command in commands:
        p = doc.add_paragraph()
        run = p.add_run(command)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)


def add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    key_run = p.add_run(f"{key}：")
    key_run.bold = True
    key_run.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(str(value))


def format_trace(trace: dict) -> str:
    retrieval = trace.get("retrievalDecision") or {}
    return (
        f"intent={trace.get('detectedIntent')}; "
        f"style={trace.get('policySummary', {}).get('style')}; "
        f"retrievalNeeded={retrieval.get('needed')}; "
        f"retrievalCalled={retrieval.get('called')}"
    )


def format_skill(skills: list[dict]) -> str:
    if not skills:
        return "无"
    first = skills[0]
    return f"{first.get('name', 'unknown')} / {first.get('status', 'unknown')} / {first.get('reason', '')}"


def format_policy(policy: dict) -> str:
    fields = [
        f"depth={policy.get('depth')}",
        f"style={policy.get('style')}",
        f"source={policy.get('source')}",
        f"useCurrentPage={policy.get('shouldUseCurrentPage')}",
        f"retrieve={policy.get('shouldRetrieveKnowledge')}",
        f"language={policy.get('answerLanguage')}",
    ]
    return "；".join(fields)


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


if __name__ == "__main__":
    main()
